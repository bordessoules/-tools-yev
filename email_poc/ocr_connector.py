"""
OCR / Text Connector — Extraction de contenu des pièces jointes
- Images : analyse via Qwen 3 VL (LM Studio) pour OCR + description
- PDF : extraction de texte via pypdf
- DOCX : extraction de texte via python-docx
Sauvegarde le résultat dans un fichier .ocr.json à côté du fichier.
"""

import base64
import json
import os
import re
import requests
from pathlib import Path

# LM Studio endpoint (Tailscale ou localhost)
LMSTUDIO_URL = "http://100.100.138.95:1234/v1/chat/completions"
MODEL_NAME = "qwen/qwen3-vl-8b"

# Extensions supportées
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
PDF_EXTENSIONS = {'.pdf'}
DOCX_EXTENSIONS = {'.doc', '.docx'}
ALL_SUPPORTED_EXTENSIONS = IMAGE_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS

# Prompt en français pour l'OCR
OCR_PROMPT = """Analyse cette image en détail. Tu dois :

1. **Décrire** ce que tu vois dans l'image (type de document, contenu visuel, objets, personnes, etc.)
2. **Transcrire** TOUT le texte visible dans l'image, en respectant la mise en page autant que possible
3. **Identifier** les éléments importants : numéros de série, références, codes-barres, dates, noms, adresses, montants, marques, modèles

Réponds en français. Sois exhaustif sur le texte — chaque mot, chaque chiffre compte.
Ne mets pas de balises markdown, juste du texte brut."""


def image_to_base64(filepath):
    """Convertir une image en base64 data URI."""
    ext = Path(filepath).suffix.lower()
    mime_map = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.webp': 'image/webp',
    }
    mime = mime_map.get(ext, 'image/jpeg')

    with open(filepath, 'rb') as f:
        data = base64.b64encode(f.read()).decode('utf-8')
    return f"data:{mime};base64,{data}"


def ocr_image(filepath, timeout=120):
    """Envoyer une image au modèle VL et récupérer la description + texte.

    Returns: dict with 'description' (str) and 'raw_response' (str)
    """
    data_uri = image_to_base64(filepath)

    payload = {
        "model": MODEL_NAME,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": data_uri}
                    },
                    {
                        "type": "text",
                        "text": OCR_PROMPT
                    }
                ]
            }
        ],
        "max_tokens": 2048,
        "temperature": 0.1,
    }

    try:
        resp = requests.post(LMSTUDIO_URL, json=payload, timeout=timeout)
        resp.raise_for_status()
        result = resp.json()
        content = result["choices"][0]["message"]["content"]
        # Strip thinking tags if present (Qwen3 sometimes wraps in <think>)
        content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
        return {
            "description": content,
            "raw_response": content,
        }
    except requests.exceptions.Timeout:
        return {"description": "[OCR timeout]", "raw_response": ""}
    except Exception as e:
        return {"description": f"[OCR error: {e}]", "raw_response": ""}


def extract_pdf_text(filepath):
    """Extraire le texte d'un PDF avec pypdf.

    Returns: dict with 'description' (str) and 'raw_response' (str)
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages_text.append(f"--- Page {i+1} ---\n{text}")

        if not pages_text:
            return {"description": "[PDF sans texte extractible]", "raw_response": ""}

        full_text = "\n\n".join(pages_text)
        num_pages = len(reader.pages)
        summary = f"Document PDF de {num_pages} page(s).\n\n{full_text}"
        return {
            "description": summary,
            "raw_response": full_text,
        }
    except Exception as e:
        return {"description": f"[PDF error: {e}]", "raw_response": ""}


def extract_docx_text(filepath):
    """Extraire le texte d'un fichier DOCX avec python-docx.

    Returns: dict with 'description' (str) and 'raw_response' (str)
    """
    try:
        from docx import Document
        doc = Document(filepath)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_texts.append(" | ".join(cells))

        if not paragraphs and not table_texts:
            return {"description": "[DOCX vide ou sans texte]", "raw_response": ""}

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n--- Tableaux ---\n" + "\n".join(table_texts)

        summary = f"Document Word ({len(paragraphs)} paragraphes).\n\n{full_text}"
        return {
            "description": summary,
            "raw_response": full_text,
        }
    except Exception as e:
        return {"description": f"[DOCX error: {e}]", "raw_response": ""}


def is_pdf(filepath):
    """Vérifie si le fichier est un PDF."""
    return Path(filepath).suffix.lower() in PDF_EXTENSIONS


def is_docx(filepath):
    """Vérifie si le fichier est un DOCX."""
    return Path(filepath).suffix.lower() in DOCX_EXTENSIONS


def is_supported(filepath):
    """Vérifie si le fichier est un type supporté (image, PDF, DOCX)."""
    return Path(filepath).suffix.lower() in ALL_SUPPORTED_EXTENSIONS


def get_ocr_cache_path(filepath):
    """Chemin du fichier cache OCR pour un fichier."""
    return filepath + ".ocr.json"


def load_ocr_cache(filepath):
    """Charger le résultat OCR depuis le cache, ou None."""
    cache = get_ocr_cache_path(filepath)
    if os.path.exists(cache):
        try:
            with open(cache, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return None
    return None


def save_ocr_cache(filepath, result):
    """Sauvegarder le résultat OCR dans le cache."""
    cache = get_ocr_cache_path(filepath)
    with open(cache, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)


def is_image(filepath):
    """Vérifie si le fichier est une image supportée."""
    return Path(filepath).suffix.lower() in IMAGE_EXTENSIONS


def process_attachments_ocr(attachments_dir, force=False):
    """Traiter toutes les pièces jointes d'un dossier (images, PDF, DOCX).

    Args:
        attachments_dir: chemin du dossier
        force: si True, re-traiter même si le cache existe

    Returns:
        dict mapping saved_name → ocr_result
    """
    results = {}
    if not os.path.isdir(attachments_dir):
        return results

    supported = [f for f in os.listdir(attachments_dir)
                 if is_supported(os.path.join(attachments_dir, f))]

    total = len(supported)
    if total == 0:
        return results

    print(f"  OCR/Texte: {total} fichiers à traiter dans {attachments_dir}")

    for i, filename in enumerate(sorted(supported)):
        filepath = os.path.join(attachments_dir, filename)

        # Vérifier le cache
        cached = load_ocr_cache(filepath)
        if cached and not force:
            results[filename] = cached
            print(f"    [{i+1}/{total}] {filename} — cache ✓")
            continue

        # Traiter selon le type
        if is_image(filepath):
            print(f"    [{i+1}/{total}] {filename} — OCR en cours...", end="", flush=True)
            result = ocr_image(filepath)
        elif is_pdf(filepath):
            print(f"    [{i+1}/{total}] {filename} — PDF extraction...", end="", flush=True)
            result = extract_pdf_text(filepath)
        elif is_docx(filepath):
            print(f"    [{i+1}/{total}] {filename} — DOCX extraction...", end="", flush=True)
            result = extract_docx_text(filepath)
        else:
            continue

        save_ocr_cache(filepath, result)
        results[filename] = result

        desc_len = len(result.get("description", ""))
        print(f" OK ({desc_len} chars)")

    return results


def load_all_ocr_results(attachments_dir):
    """Charger tous les résultats OCR/texte disponibles (depuis le cache uniquement).

    Returns: dict mapping saved_name → ocr_result
    """
    results = {}
    if not os.path.isdir(attachments_dir):
        return results

    for filename in os.listdir(attachments_dir):
        filepath = os.path.join(attachments_dir, filename)
        if not is_supported(filepath):
            continue
        cached = load_ocr_cache(filepath)
        if cached:
            results[filename] = cached

    return results


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        # Process a single image
        filepath = sys.argv[1]
        if os.path.isfile(filepath):
            print(f"OCR: {filepath}")
            result = ocr_image(filepath)
            save_ocr_cache(filepath, result)
            print(result["description"])
        elif os.path.isdir(filepath):
            process_attachments_ocr(filepath, force="--force" in sys.argv)
    else:
        # Default: process all client attachments
        att_base = os.path.join(os.path.dirname(__file__), "attachments")
        if os.path.isdir(att_base):
            for client_dir in os.listdir(att_base):
                client_path = os.path.join(att_base, client_dir)
                if os.path.isdir(client_path):
                    print(f"\nClient: {client_dir}")
                    process_attachments_ocr(client_path)
